from docx import Document
from docx.shared import Inches
import pyttsx3


def speak(text):
    pyttsx3.speak(text)


document = Document()

# add profile image
document.add_picture('teri.png', width=Inches(2.0))

# Detail information of the user
name = input('What is your name? ')
# call the pyttsx3
speak(f'Hello {name}, how are you today? ')
phone_number = input('What is your phone number? ')
email = input('What is your email? ')

document.add_paragraph(
    f'{name} | {phone_number} | {email}'
)

# about me
document.add_heading('About me')
about_me = input('Why are you special amongst others? ')
document.add_paragraph(about_me)

# Work experience
document.add_heading('Work experience')
p = document.add_paragraph()

company = input('Enter company ')
from_date = input('From Date ')
to_date = input('To Date ')

p.add_run(f'{company} ').bold = True
p.add_run(f'{from_date} - {to_date} \n').italic = True

experience_details = input(f'Describe your experience at {company} ')
p.add_run(experience_details)

# more experiences
while True:
    has_more_experiences = input('Do you have more experiences? Yes or No ')
    if has_more_experiences.lower() == 'yes' or has_more_experiences.lower() == 'y':
        p = document.add_paragraph()
        company = input('Enter company ')
        from_date = input('From Date ')
        to_date = input('To Date ')

        p.add_run(f'{company} ').bold = True
        p.add_run(f'{from_date} - {to_date} \n').italic = True

        experience_details = input(f'Describe your experience at {company} ')
        p.add_run(experience_details)
    else:
        break

# showcase skills
document.add_heading('Skills')
skill = input('Enter skill ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have more skills? Yes or No ')
    if has_more_skills.lower() == 'yes' or has_more_skills.lower() == 'y':
        skill = input('Enter skill ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break


# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'CV generated using Python scripts and developed by Teri'


document.save('cv.docx')
